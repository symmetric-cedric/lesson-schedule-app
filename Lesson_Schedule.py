import streamlit as st
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# Display Logo (uncomment and set path if needed)
st.image("logo.png", width=200)


# Function Definitions

def generate_schedule(total_lessons, frequency_days, start_date):
    freq_idxs = set(weekday_map[d] for d in frequency_days)
    lessons, skipped = [], []
    date = start_date
    while len(lessons) < total_lessons:
        if date.weekday() in freq_idxs:
            if date in holiday_dates:
                skipped.append(date)
            else:
                lessons.append(date)
        date += timedelta(days=1)
    return lessons, skipped


def calculate_week_range(total_lessons, freq_per_week, lesson_dates):
    if total_lessons in (10, 30):
        return total_lessons
    key = freq_per_week if freq_per_week < 3 else 3
    m = {
        1: {4:5,12:15,24:30},
        2: {8:5,24:15,48:30},
        3: {12:5,36:15,72:30}
    }
    base = m.get(key, {}).get(total_lessons, 5)
    holidays = sum(1 for d in lesson_dates if d in holiday_dates)
    return base + holidays


def calculate_main_course_fee(lessons_per_week, total_lessons):
    pricing = {
        (1, 4): (1280, 50), (1,12): (3456,100), (1,24): (6144,150),
        (2,8): (2400,100), (2,24):(5760,150),(2,48):(10080,300),
        (3,12):(3360,100),(3,36):(7560,250),(3,72):(14112,400),
        (None,10):(3500,100),(None,30):(9000,150)
    }
    return pricing.get((lessons_per_week, total_lessons)) or pricing.get((None,total_lessons),(0,0))



def calculate_value_added_fee(total_lessons, value_added_courses):
    if not value_added_courses:
        return 0

    fee = 0
    for _ in value_added_courses:
        if total_lessons in (4, 8):
            fee += 100 * total_lessons
        elif total_lessons == 12:
            fee += 75 * total_lessons
        elif total_lessons >= 24:
            fee += 50 * total_lessons
        else:
            fee += 0  # Default or unsupported lesson count
    return fee



def calculate_optional_items(selected):
    fee, details = 0, []
    for opt in selected:
        amt = None
        if opt in optional_items_map:
            amt = optional_items_map[opt]
        elif "（＋$" in opt:
            amt = int(opt.split("（＋$")[-1].replace("）",""))
        if amt is not None:
            fee += amt
            details.append((opt, amt))
    return fee, details


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")  # Create a new paragraph element
    paragraph._p.addnext(new_p)  # Insert after the current paragraph

    # Add a run (text span) to the new paragraph
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)

    # Wrap it in a Paragraph object so you can keep chaining if needed
    return Paragraph(new_p, paragraph._parent)

def fill_template_doc(
    student_name, branch_name, invoice_number,
    main_tuition, main_material,
    value_tuition, value_material,
    optional_items,
    start_date, lesson_dates, week_range,
    day_time_pairs, skipped_holidays,
    template_path,
    subjects, value_added_courses,value_material_selections
):
    doc = Document(template_path)
    
    bank_info_map = {
    "九龍灣(淘大)分校": "恆生銀行\n賬戶名稱：YCH EDUCATION LIMITED\n賬戶號碼：369 439 963 883\n快速支付編號 : 161107016",
    "藍田(麗港城)分校": "恆生銀行\n賬戶名稱：YCL EDUCATION LIMITED\n賬戶號碼：244 257 796 883\n快速支付編號 : 161258090",
    "青衣(青怡)分校": "匯豐銀行\n賬戶名稱：LCH EDUCATION LIMITED\n賬戶號碼：049 656 150 838\n快速支付編號 : 160832457",
    "九龍站(港景峯)分校": "恆生銀行\n賬戶名稱：MISS MAN EDUCATION LIMITED\n賬戶號碼：244 333 787 883\n快速支付編號 : 103 779 914",
    "鑽石山(萬迪廣場)分校": "恆生銀行\n賬戶名稱：MISS MAN EDUCATION LIMITED\n賬戶號碼：244 333 787 883\n快速支付編號 : 103 779 914"
    }

    main_subjects_str = '、'.join(subjects)
    value_added_courses_str = '、'.join(value_added_courses)

    
    # Replace placeholders
    start_date_str = start_date.strftime('%d/%m/%Y')
    end_date = start_date + timedelta(weeks=week_range) - timedelta(days=1)
    date_range_str = f"{start_date_str} 至 {end_date.strftime('%d/%m/%Y')}"

    # Construct the 上課時間 string
    day_time_str = ' / '.join(f"{day} {time}" for day, time in day_time_pairs.items())

    replacements = {
        "單號:": f"單號: {invoice_number}",
        "學生姓名：": f"學生姓名：{student_name}",
        "堂數": f"堂數：{total_lessons}",
        "上課期數範圍": f"上課期數範圍：{date_range_str}",
        "分校": branch_name,
        "上課時間：": f"上課時間：{day_time_str}",
        "Bank_info": bank_info_map[branch_name],
    }

    for para in doc.paragraphs:
        for key, new_text in replacements.items():
            if para.text.strip().startswith(key):
                para.text = new_text

    # Find "學費計算" to insert after
    fee_idx = next((i for i, p in enumerate(doc.paragraphs) if "學費計算" in p.text), None)
    if fee_idx is not None:
        base_para = doc.paragraphs[fee_idx]
    
        current_para = insert_paragraph_after(base_para, f"主科（{main_subjects_str}）：+${main_tuition}")
        current_para = insert_paragraph_after(current_para, f"小組活動教材：+${main_material}")
        current_para = insert_paragraph_after(current_para, f"增值課程（{value_added_courses_str}）：+${value_tuition}")

        # 增值課程教材 breakdown
        value_material_total = 0
        if value_material_selections:
            current_para = insert_paragraph_after(current_para, "增值課程教材：")
            for course, lesson_count in value_material_selections.items():
                try:
                    price = value_material[course][lesson_count]
                    current_para = insert_paragraph_after(current_para, f"{course}（{lesson_count}）：+${price}")
                    value_material_total += price
                except KeyError:
                    current_para = insert_paragraph_after(current_para, f"{course}（{lesson_count}）：資料錯誤，請檢查設定")
        
        # 其他項目
        current_para = insert_paragraph_after(current_para, "其他:")
        for opt, amt in optional_items:
            current_para = insert_paragraph_after(current_para, f"{opt}：{'+' if amt > 0 else ''}${amt}")
        
        # 計算總額
        total_amount = main_fee + main_material + value_fee + value_material_total + opt_fee
        if total_lessons == (24 or 36):
            insert_paragraph_after(current_para, f"總額：= ${total_amount} (銀行轉賬) / ${total_amount-50}(現金到校繳付)")
        elif total_lessons == (48 or 72):
            insert_paragraph_after(current_para, f"總額：= ${total_amount} (銀行轉賬) / ${total_amount-100}(現金到校繳付)")
        else:
            insert_paragraph_after(current_para, f"總額：= ${total_amount} (銀行轉賬)")


    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf







# Weekday and Holiday Setup
weekday_map = {
    "星期一": 0, "星期二": 1, "星期三": 2, "星期四": 3,
    "星期五": 4, "星期六": 5, "星期日": 6
}
weekday_chinese = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']

public_holidays = {
    "1 July 2026", "26 September 2026", "1 October 2026", "19 October 2026",
    "25 December 2026", "26 December 2026", 
    "1 January 2027", "6 February 2027", "8 February 2027", "9 February 2027",
    "26 March 2027", "27 March 2027", "29 March 2027", "5 April 2027",
    "1 May 2027", "13 May 2027", "9 June 2027", "1 July 2027",
    "16 September 2027", "1 October 2027", "8 October 2027",
    "25 December 2027", "27 December 2027"
}


holiday_dates = set(datetime.strptime(d, "%d %B %Y").date() for d in public_holidays)

template_path = "Testing.docx"

lesson_time_options = [
    "9:30-11:00", "10:00-11:30", "10:30-12:00", "11:00-12:30",
    "11:30-13:00", "12:00-13:30", "13:30-15:00", "14:00-15:30",
    "14:30-16:00", "15:00-16:30", "15:30-17:00", "16:00-17:30",
    "16:30-18:00", "17:00-18:30", "17:30-19:00"
]

subject_options = [
    "中文記憶閱讀", "英文拼音", "小一面試班", "小學銜接班", "小學精進班"
]

value_added_options = [
    "英文拼音", "高效寫字", "聆聽訓練", "說話訓練",
    "思維閱讀", "創意理解", "作文教學"
]

# Optional items configuration

value_material = {
    "英文拼音課本": {
        "單本": 50
    },
    "中文認字課本": {
        "單本": 50
    },
    "高效寫字課本": {
        "單本": 50
    },
    "創意理解課本": {
        "單本": 50
    },
    "創意理解・語文工作紙": {
        "4堂": 50,
        "8堂": 100,
        "12堂": 100,
        "24堂": 150,
        "36堂": 250,
        "48堂": 300,
        "72堂": 400,
    },
    "聆聽訓練教材": {
        "4堂": 50,
        "8堂": 100,
        "12堂": 100,
        "24堂": 150,
        "36堂": 250,
        "48堂": 300,
        "72堂": 400,
    },
    "高效寫字教材": {
        "4堂": 50,
        "8堂": 100,
        "12堂": 100,
        "24堂": 150,
        "36堂": 250,
        "48堂": 300,
        "72堂": 400,
    },
    "思維閱讀教材": {
        "4堂": 50,
        "8堂": 100,
        "12堂": 100,
        "24堂": 150,
        "36堂": 250,
        "48堂": 300,
        "72堂": 400,
    },
    "作文教學工作紙": {
        "4堂": 50,
        "8堂": 100,
        "12堂": 100,
        "24堂": 150,
        "36堂": 250,
        "48堂": 300,
        "72堂": 400,
    },
}


optional_items_map = {
    "試堂日報讀贈券：即日報讀可獲舊生推薦現金券 ($100)": -100,
    "試堂日報讀贈券：即日報讀可扣減試堂費 ($200)": -200,
    "在學證明": -50
}


# Streamlit UI
st.title(":calendar: 課程收據單生成器")

# User Inputs
student_name = st.text_input("學生姓名")
branch_name = st.selectbox("分校名稱", [
    "九龍灣(淘大)分校", "藍田(麗港城)分校", "青衣(青怡)分校",
    "九龍站(港景峯)分校", "鑽石山(萬迪廣場)分校"
])
invoice_number = st.text_input("單號")
total_lessons = st.selectbox("堂數", [4, 8, 10, 12, 24, 30, 36, 48, 72])

st.subheader("上課日及時間")
day_time_pairs = {}
for day in weekday_map:
    if st.checkbox(day):
        day_time_pairs[day] = st.selectbox(f"{day} 上課時間", lesson_time_options, key=day)

subjects = st.multiselect("主科", subject_options)
value_added_courses = st.multiselect("增值課程", value_added_options)
start_date = st.date_input("開始日期")

# --- PREVIEW lesson dates for cancellation UI ---
day_names_selected = list(day_time_pairs.keys())
preview_lesson_dates, _ = generate_schedule(
        total_lessons, day_names_selected, start_date
    )
show_cancel = st.checkbox("是否有取消上課日期？", value=False)
st.markdown("---")
st.markdown("<hr style='border: 1px solid #ccc;'>", unsafe_allow_html=True)
st.markdown("### 🧾 學費資訊")

cancel_holidays = []

if show_cancel:
    cancel_holidays = st.multiselect(
        "取消上課日期",
        options=preview_lesson_dates,
        format_func=lambda d: d.strftime('%Y/%m/%d（%A）')
    )

holiday_dates.update(cancel_holidays)


# UI: value-added materials selection with lesson count
value_material_selections = {}
for course in value_material:
    if st.checkbox(course):
        lesson_option = st.selectbox(
            f"{course} - 選擇堂數", list(value_material[course].keys()), key=course
        )
        value_material_selections[course] = lesson_option

# Use the defined map for optional selections
optional_selections = st.multiselect("其他選項", list(optional_items_map.keys()))






# Generate Receipt
if st.button("生成收據單"):
    # Validate
    if not all([student_name, branch_name, invoice_number, subjects, day_time_pairs]):
        st.error("請填妥所有必填欄位。")
    else:
        # Compute schedules and fees
        lesson_dates, skipped_holidays = generate_schedule(
            total_lessons, list(day_time_pairs.keys()), start_date
        )
        
        week_range = calculate_week_range(
            total_lessons, len(day_time_pairs), lesson_dates
        )
        
        # Fees
        main_fee, main_material = calculate_main_course_fee(len(day_time_pairs), total_lessons)
        value_fee = calculate_value_added_fee(total_lessons, value_added_courses)
        # assume no separate materials for value-added or adjust as needed
        opt_fee, opt_details = calculate_optional_items(optional_selections)






        
        # Fill and download document
        doc_file = fill_template_doc(
            student_name, branch_name, invoice_number,
            main_fee, main_material,
            value_fee, value_material,
            opt_details,
            start_date, lesson_dates, week_range,
            day_time_pairs, skipped_holidays,
            template_path, subjects, value_added_courses,value_material_selections
        )
        st.success("收據單已生成！")
        st.download_button(
            "下載 Word 文件", data=doc_file,
            file_name="課程收據單.docx"
        )











        
        # Clipboard Text
        end_date = start_date + timedelta(weeks=week_range) - timedelta(days=1)

        bill_text_lines = [
            f"分校：{branch_name}",
            f"單號：{invoice_number}",
            f"學生姓名：{student_name}",
            f"堂數：{total_lessons}",
            f"主科：{' / '.join(subjects)}",
            f"增值課程：{' / '.join(value_added_courses)}",
            f"📆 上課期數範圍：{start_date.strftime('%d/%m/%Y')} 至 {end_date.strftime('%d/%m/%Y')}",
            f"📌 下期學費繳交日期：{lesson_dates[-1].strftime('%d/%m/%Y')}",
            "",
            "上課時間："
        ] + [f"{day} {time}" for day, time in day_time_pairs.items()] + [
            "",
            "📅 上課日期安排："
        ] + [
            f"{i}. {d.strftime('%d/%m/%Y')} ({weekday_chinese[d.weekday()]})"
            for i, d in enumerate(lesson_dates, 1)
        ]
        
        if skipped_holidays:
            bill_text_lines += ["", "❌ 公眾假期 (休息) / 取消上課日期:"] + [
                f"- {d.strftime('%d/%m/%Y')} ({weekday_chinese[d.weekday()]})" for d in skipped_holidays
            ]
        
        # Additional Notices
        bill_text_lines += [
            "",
            "補堂：",
            "補堂時間一經確定，不可更改。缺席補堂將不會再安排補堂",
            "\n➿➿➿➿➿➿➿➿➿➿",
            "📣家長須知 📣",
            "1. 返學安排🎒",
            "- 上課前，須先上洗手間🚾",
            "- 學生遲到或無故缺席均不設補時或補課❌",
            "2. 放學安排",
            "- 家長須準時接送子女放學 ⏰",
            "❌ 3. 公眾假期 (休息):" if skipped_holidays else "❌ 3. 公眾假期 (休息)",
        ]
        
        if skipped_holidays:
            bill_text_lines += [
                f"- {d.strftime('%d/%m/%Y')} ({weekday_chinese[d.weekday()]})" for d in skipped_holidays
            ]
        
        bill_text_lines += [
            "4. 請假安排：",
            "i.事假",
            "- 須上課3天前以短訊通知，方可安排補堂",
            "- 補堂須於課程結束日前完成",
            "- 不足3天或即日通知，不設❌補堂❌",
            "ii.病假😷",
            "- 須後補醫生証明📝，方可安排補堂",
            "- 補堂須於課程結束日前完成",
            "5. 惡劣天氣安排：",
            "天文台於上課前兩小時發出惡劣天氣警告信號，本中心將作出以下安排：",
            "- 黃色、紅色暴雨警告",
            "- ⁠三號颱風訊號",
            "✅本中心會照常上課✅（家長可自行決定子女會否上課，上課前2小時以短訊通知請假，可安排補堂）"
        ]
        
        # Final output
        bill_text = '\n'.join(bill_text_lines)


        
        st.subheader("📋 複製以下文字：")
        st.code(bill_text, language="text")




